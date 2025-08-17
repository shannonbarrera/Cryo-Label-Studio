import sys
import re
import copy
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from datetime import datetime, date
from label_templates import label_templates
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import math



def get_row_and_column_indices(templatepath, table_format):
    labelsheet = Document(templatepath)
    table = labelsheet.tables[0]
    if table_format == "checkerboard":
        row_indices = [i for i in range(len(table.rows)) if i % 2 == 0]
        col_indices = [j for j in range(len(table.columns)) if j % 2 == 0]

    if table_format == "LSL stripes":
        row_indices = [i for i in range(len(table.rows))]
        col_indices = [j for j in range(len(table.columns)) if j % 2 == 0]

    return row_indices, col_indices


def get_first_page_row_indices(start_row, end_row, row_indices):
    first_page_row_indices = []
    for i in range(len(row_indices)):
        if i >= start_row - 1 and i <= end_row - 1:
            first_page_row_indices.append(row_indices[i])
    return first_page_row_indices


def get_first_page_col_indices(start_col, end_col, start_row, end_row, col_indices):
    first_page_first_row_col_indices = []
    first_page_last_row_col_indices = []
    if end_row - start_row == 0:
        for i in range(len(col_indices)):
            if i >= start_col - 1 and i <= end_col - 1:
                first_page_first_row_col_indices.append(col_indices[i])
        return first_page_first_row_col_indices, first_page_last_row_col_indices

    for i in range(len(col_indices)):
        if i >= start_col - 1:
            first_page_first_row_col_indices.append(col_indices[i])

    for i in range(len(col_indices)):
        if i <= end_col - 1:
            first_page_last_row_col_indices.append(col_indices[i])

    if end_col - start_col == 0:
        first_page_first_row_col_indices = []
        for i in range(len(col_indices)):
            if i >= start_col - 1:
                first_page_first_row_col_indices.append(col_indices[i])

    return first_page_first_row_col_indices, first_page_last_row_col_indices


def get_max_labels_per_page(spec, templatepath, table_format):
    """
    Calculates how many label entries can fit per page.

    Args:
        labeltemplate (str): Path to label template file.
        labelsheetlayouttype (str): Layout type of the label sheet.
        copiesperlabel (int): Number of repeated labels per entry.

    Returns:
        int: Maximum number of unique label entries per page.
    """

    row_indices, column_indices = get_row_and_column_indices(templatepath, table_format)
    total_cells = len(row_indices) * len(column_indices)
    return total_cells


def get_max_labels_first_page(
    first_page_row_indices,
    column_indices,
    first_page_first_row_col_indices,
    first_page_last_row_col_indices,
):
    labels_first_page_first_last_row = len(first_page_first_row_col_indices) + len(
        first_page_last_row_col_indices
    )
    labels_first_page_middle_rows = (len(first_page_row_indices) - len(first_page_row_indices)) * len(
        column_indices
    )
    total_cells = labels_first_page_first_last_row + labels_first_page_middle_rows
    return total_cells


def paginate_labels(
    first_page_max_labels, max_labels_per_page, data_list, copiesperlabel
):
    total_labels = len(data_list) * copiesperlabel
    if first_page_max_labels > total_labels:
        num_pages = 1
        firstpage = []
        pages = []
        for item in data_list:
            for _ in range(copiesperlabel):
                firstpage.append(item)

    else:
        num_pages = (
            math.ceil((total_labels - first_page_max_labels) / max_labels_per_page) + 1
        )
        firstpage = []

        first_page_max_data = first_page_max_labels // copiesperlabel

        if first_page_max_data * copiesperlabel < first_page_max_labels:
            first_page_max_data += 1

        for item in data_list[:first_page_max_data]:
            for i in range(copiesperlabel):
                firstpage.append(item)
        if len(firstpage) > first_page_max_labels:
            hangover = firstpage[first_page_max_labels:]
        else:
            hangover = []
   
        firstpage = firstpage[:first_page_max_labels]
        remaining = data_list[first_page_max_data:]
        pages = []
        remainingindex = 0
        page = []
        for i in range(1, num_pages):
            if remainingindex == 0:
                page = hangover
            remaining_labels_on_page = max_labels_per_page - len(page)
            remaining_data_items_on_page = math.ceil(
                remaining_labels_on_page / copiesperlabel
            )
            for item in remaining[remainingindex:remaining_data_items_on_page]:
                for i in range(copiesperlabel):
                    page.append(item)
                remainingindex += 1

            if len(page) > max_labels_per_page:
                hangover = page[max_labels_per_page:]
            else:
                hangover = []
            page = page[:max_labels_per_page]
            pages.append(page)

            remaining = remaining[remainingindex:]
            remainingindex = 0
    return firstpage, pages


def format_labels_page(
    data_list,
    templatepath,
    first_page_row_indices,
    column_indices,
    first_page_first_row_col_indices,
    first_page_last_row_col_indices,
    spec,
    needs_page_break,
    is_last_page=False
):
    labelsheet = Document(templatepath)
    table = labelsheet.tables[0]
    textboxformatinput = spec.textboxformatinput
    fontname = spec.fontname
    fontsize = spec.fontsize

    first_row = first_page_row_indices[0]

    if len(first_page_row_indices) > 2:
        middle_rows = first_page_row_indices[1:-1]
        last_row = first_page_row_indices[-1]
    else:
        middle_rows = []
        last_row = first_page_row_indices[-1]

    labelcount = 0
    # Fill first row
    for cind in first_page_first_row_col_indices:
        if labelcount >= len(data_list):
            break
        current_cell = table.rows[first_row].cells[cind]
        format_label_cell(
            current_cell,
            data_list[labelcount],
            textboxformatinput,
            fontname,
            fontsize,
            spec.alignment,
            spec.date_format,
            spec.identical_or_incremental,
        )
        labelcount += 1

    # Fill middle rows
    for row in middle_rows:
        for cind in column_indices:
            if labelcount >= len(data_list):
                break
            current_cell = table.rows[row].cells[cind]
            format_label_cell(
                current_cell,
                data_list[labelcount],
                textboxformatinput,
                fontname,
                fontsize,
                spec.alignment,
                spec.date_format,
                spec.identical_or_incremental,
            )
            labelcount += 1

    # Fill last row (if it’s different)
    if last_row != first_row:
        for cind in first_page_last_row_col_indices:
            if labelcount >= len(data_list):
                break
            current_cell = table.rows[last_row].cells[cind]
            format_label_cell(
                current_cell,
                data_list[labelcount],
                textboxformatinput,
                fontname,
                fontsize,
                spec.alignment,
                spec.date_format,
                spec.identical_or_incremental,
            )
            labelcount += 1

    if needs_page_break:
        if not is_last_page:
            labelsheet.add_page_break()

    return labelsheet


def format_label_cell(cell, data, textboxformatinput, fontname, fontsize, alignment, date_format, identical_or_incremental=None):
    """
    Populates a single label cell with the given data and formats it according to the label template.

    This function takes the extracted data and fills the corresponding label cell with the entry
    number, last name, first name, and date.

    Args:
        cell (docx.table._Cell): The label cell in which the data will be inserted.
        data (list): A list containing the entry's label data. The expected structure is:
                 [identifier, last name, first name, date]. The `date` is expected to be a `datetime` object, but the program will also work if it's a str.

    Returns:
        None: The function directly modifies the cell's content and style.

    Example:
        format_label_cell(cell, ["25-0001", "Doe", "John", datetime.datetime(2025, 3, 29)])
        This would fill the cell with the following formatted text:
        "25-0001
         Doe, John
         03/29/2025"
    """
    if identical_or_incremental != "Identical":
        if textboxformatinput:
            label_text = apply_format_to_row(textboxformatinput, data, date_format)
            cell.text = label_text

        else:
            cell.text = data
    else:
        cell.text = data

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment_map.get(
            alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER
        )
        for run in paragraph.runs:
            run.font.size = Pt(float(fontsize))
            run.font.name = fontname
            run.bold = True
    return

import re

def smart_wrap_label_text(label_text, max_chars, prefix=None, buffer=3):
    """
    Inserts a line break after the prefix if the label is close to overflowing.

    Args:
        label_text (str): The full label string
        max_chars (int): Max characters allowed per line
        prefix (str, optional): The known prefix (will use len(prefix))
        prefix_len (int, optional): If no prefix string, supply character length directly
        buffer (int): Number of chars before max to trigger wrapping

    Returns:
        str: Wrapped label text with newline inserted after prefix (if needed)
    """
    if len(label_text) <= max_chars - buffer:
        return label_text


    if prefix is not None:
        prefix_len = len(prefix)
        # Insert \n after prefix
        return label_text[:prefix_len] + "\n" + label_text[prefix_len:]
    
    # Otherwise split at last space before limit
    last_space = label_text.rfind(" ", 0, max_chars - buffer)
    if last_space != -1:
        return label_text[:last_space] + "\n" + label_text[last_space+1:]

    # Fallback: force break
    return label_text[:max_chars - buffer] + "\n" + label_text[max_chars - buffer:]



def parse_slice(slice_str):
    """
    Safely parses a slice string like [5:], [:10], [2:5] into a Python slice object.
    """
    match = re.match(r"\[(?:(\d*))?:(?:(\d*))?\]", slice_str)
    if not match:
        raise ValueError(f"Invalid slice format: {slice_str}")
    start_str, end_str = match.groups()
    start = int(start_str) if start_str else None
    end = int(end_str) if end_str else None
    return slice(start, end)

def apply_format_to_row(textboxformatinput, row_data, date_format):
    """
    Applies a label format string with placeholders to a row of data.
    Supports optional slicing like {FIELD}[2:], skips None values, formats dates.

    Args:
        textboxformatinput (str): A string with placeholders like "{SampleID}\n{Date}".
        row_data (list): A list of values in the same order as placeholders.

    Returns:
        str: The formatted label string.
    """
    slice_pattern = re.compile(r"({([^}]+)}(\[[^\]]+\])?)")
    placeholder_to_value = {}
    matches = list(slice_pattern.finditer(textboxformatinput))

    for i, match in enumerate(matches):
        full_placeholder = match.group(1)  # e.g., {SERUM ID}[6:]
        key = match.group(2)               # e.g., SERUM ID
        slice_part = match.group(3)        # e.g., [6:]

        if i >= len(row_data):
            placeholder_to_value[full_placeholder] = ""
        else:
            value = row_data[i]

            if isinstance(value, (datetime, date)):
                if date_format == "Leave as is":
                    value = str(value)  # Use raw datetime string (likely from Excel)
                elif date_format:
                    value = value.strftime(date_format)
                else:
                    value = str(value)

            value = "" if value is None else str(value)

            if slice_part:
                try:
                    slice_obj = parse_slice(slice_part)
                    value = value[slice_obj]
                except Exception as e:
                    print(f"Warning: invalid slice {slice_part} on {key}: {e}")

            placeholder_to_value[full_placeholder] = value

    result = textboxformatinput
    for full_placeholder, value in placeholder_to_value.items():
        result = result.replace(full_placeholder, value)

    return result

def combine_docs(doc1, doc2):
    """
    Appends all content from doc2 into doc1.

    Args:
        doc1 (Document): First document.
        doc2 (Document): Second document to append.

    Returns:
        Document: Combined document.
    """
    for element in doc2.element.body:
        doc1.element.body.append(element)

    return doc1
